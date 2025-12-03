import streamlit as st
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import io
import PyPDF2
from PIL import Image
import pytesseract

# Set path Tesseract untuk Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Konfigurasi API OpenAI (AMAN melalui secrets)
OPENAI_API_KEY = st.secrets["openai"]["api_key"]
OPENAI_MODEL = "gpt-5-mini-2025-08-07"  # atau "gpt-4o-mini" jika ingin lebih hemat

# Konfigurasi halaman
st.set_page_config(
    page_title="Rapport Writer Assistance",
    page_icon="📊",
    layout="wide"
)

# Fungsi untuk normalisasi nama HSH
def normalize_hsh(hsh_name):
    if pd.isna(hsh_name):
        return ""
    normalized = str(hsh_name).strip().upper()
    normalized = ' '.join(normalized.split())
    return normalized

def find_matching_hsh(target_hsh, hsh_list):
    target_normalized = normalize_hsh(target_hsh)
    for hsh in hsh_list:
        if normalize_hsh(hsh) == target_normalized:
            return hsh
    for hsh in hsh_list:
        normalized = normalize_hsh(hsh)
        if target_normalized in normalized or normalized in target_normalized:
            return hsh
    return None

@st.cache_data
def load_excel_files():
    try:
        skor_total = pd.read_excel('documents/SKOR_TOTAL_ALL.xlsx', sheet_name='SKOR TOTAL_ALL')
        skor_survei = pd.read_excel('documents/Skor_SURVEI_ALL.xlsx', sheet_name='Skor_SURVEI_ALL_FUNGSI')
        skor_benchmark_evidence = pd.read_excel('documents/Skor_benchmark.xlsx', sheet_name='Evidence')
        skor_benchmark_survei = pd.read_excel('documents/Skor_benchmark.xlsx', sheet_name='Survei')
        
        if 'HSH' in skor_total.columns:
            skor_total['HSH_normalized'] = skor_total['HSH'].apply(normalize_hsh)
        if 'HSH' in skor_survei.columns:
            skor_survei['HSH_normalized'] = skor_survei['HSH'].apply(normalize_hsh)
        
        skor_benchmark_evidence['HSH_normalized'] = skor_benchmark_evidence.iloc[:, 0].apply(normalize_hsh)
        skor_benchmark_survei['HSH_normalized'] = skor_benchmark_survei.iloc[:, 0].apply(normalize_hsh)
        
        return skor_total, skor_survei, skor_benchmark_evidence, skor_benchmark_survei
    except Exception as e:
        st.error(f"Error loading Excel files: {str(e)}")
        st.info("Pastikan folder 'documents' ada dan berisi file: SKOR_TOTAL_ALL.xlsx, Skor_SURVEI_ALL.xlsx, dan Skor_benchmark.xlsx")
        return None, None, None, None

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_image(image_file):
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image, lang='ind+eng')
        return text
    except Exception as e:
        return f"Error reading image: {str(e)}"

def read_uploaded_file(uploaded_file):
    if uploaded_file is None:
        return None
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_extension in ['xlsx', 'xls']:
            df = pd.read_excel(uploaded_file)
            return df.to_string()
        elif file_extension == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_extension in ['png', 'jpg', 'jpeg']:
            return extract_text_from_image(uploaded_file)
        else:
            return "Format file tidak didukung"
    except Exception as e:
        return f"Error reading file: {str(e)}"

# Fungsi untuk memanggil OpenAI API
def call_openai(prompt, max_tokens=4000):
    """Memanggil OpenAI API untuk analisis"""
    try:
        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": OPENAI_MODEL,
            "messages": [
                {
                    "role": "system",
                    "content": """Anda adalah konsultan senior budaya kerja perusahaan yang berpengalaman dengan pendekatan apresiatif dan profesional. 

TONE & GAYA KOMUNIKASI:
- Gunakan bahasa yang apresiatif, menghargai usaha yang telah dilakukan
- Profesional namun hangat dan mendukung
- Fokus pada kekuatan (strength-based approach) sebelum memberikan saran perbaikan
- Hindari kata-kata negatif atau menghakimi
- Gunakan frasa seperti "telah menunjukkan komitmen yang baik", "dapat lebih dioptimalkan", "peluang untuk pengembangan lebih lanjut"
- Berikan apresiasi spesifik terhadap pencapaian yang ada

FOKUS ANALISIS:
- Fokus pada aspek PERILAKU (behavior): perubahan mindset, kolaborasi, komunikasi, kepemimpinan, keterlibatan, partisipasi
- Hindari aspek teknis operasional
- Berikan analisis yang singkat, padat, jelas, dan actionable
- Setiap poin harus spesifik dan dapat ditindaklanjuti

FORMAT OUTPUT:
- Mulai dengan apresiasi umum
- "Hal yang Sudah Baik" harus spesifik dan menghargai pencapaian
- "Hal yang Dapat Diperbaiki" disampaikan sebagai peluang pengembangan, bukan kritik"""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.7,
            "max_tokens": max_tokens
        }
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data, timeout=60)
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content']
        else:
            return f"Error calling OpenAI API: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Exception in OpenAI API call: {str(e)}"

# === Fungsi Analisis (semua menggunakan call_openai) ===

def analyze_strategi_budaya(pcb_content):
    prompt = f"""
Analisis form PCB berikut dengan pendekatan APRESIATIF dan PROFESIONAL, fokus pada aspek PERILAKU:

{pcb_content}

EVALUASI:
1. Apakah Goals/Business Initiatives/Improvement menggunakan metode SMART (Specific, Measurable, Achievable, Relevant, Time-bound)?
2. Apakah ada kerunutan logis dari identifikasi kendala/Peluang Perbaikan Bisnis ke Business Initiatives/improvement?
3. Apakah PCB lengkap dan utuh dalam menggambarkan strategi budaya?

FOKUS UTAMA: Aspek PERILAKU seperti:
- Perubahan mindset dan pola pikir
- Peningkatan kolaborasi antar tim
- Perbaikan komunikasi internal
- Penguatan kepemimpinan
- Peningkatan keterlibatan dan partisipasi pekerja
- Penerapan nilai-nilai AKHLAK dalam keseharian

TONE: Apresiatif, profesional, dan konstruktif

Berikan output dalam format:

**Apresiasi Umum:**
[Berikan apresiasi terhadap upaya dan komitmen yang telah ditunjukkan dalam penyusunan PCB, fokus pada aspek positif yang terlihat]

**Hal yang Sudah Baik:**
- [Poin spesifik 1 - apresiasi pencapaian konkret terkait perilaku]
- [Poin spesifik 2 - apresiasi pencapaian konkret terkait perilaku]
- [Poin spesifik 3 - jika ada]

**Peluang Pengembangan Lebih Lanjut:**
- [Saran 1 - disampaikan sebagai peluang, bukan kritik, fokus perilaku]
- [Saran 2 - disampaikan sebagai peluang, bukan kritik, fokus perilaku]
- [Saran 3 - jika perlu]
"""
    return call_openai(prompt)

def analyze_program_budaya(pcb_content):
    prompt = f"""
Analisis Program Budaya dari form PCB berikut dengan pendekatan APRESIATIF dan PROFESIONAL, fokus pada aspek PERILAKU:

{pcb_content}

EVALUASI PROGRAM:
1. **Program Standar (One Hour Meeting)**: Kualitas dialog, keterbukaan komunikasi, partisipasi aktif
2. **Program Mandatory (ONE Action)**: Implementasi aksi nyata, keterlibatan pekerja, dampak perilaku
3. **Program Spesifik (ONE KOLAB)**: Kolaborasi lintas fungsi, sinergi tim, inovasi bersama

Untuk setiap program, evaluasi:
- Kesesuaian judul dengan tujuan perubahan perilaku
- Kualitas deliverables dalam mendorong perubahan perilaku
- Kontribusi program terhadap pencapaian Goals/Business Initiatives
- Tingkat keterlibatan dan partisipasi pekerja

FOKUS: Aspek PERILAKU (komunikasi, kolaborasi, keterlibatan, perubahan mindset)

TONE: Apresiatif, profesional, dan mendukung

Berikan output dalam format:

**Apresiasi Umum:**
[Apresiasi terhadap desain dan implementasi program budaya yang telah dilakukan, soroti komitmen tim]

**Hal yang Sudah Baik:**
- [Apresiasi spesifik program 1 - fokus dampak perilaku positif]
- [Apresiasi spesifik program 2 - fokus dampak perilaku positif]
- [Apresiasi spesifik program 3 - jika ada]

**Peluang Pengembangan Lebih Lanjut:**
- [Saran pengembangan 1 - sebagai peluang optimalisasi, fokus perilaku]
- [Saran pengembangan 2 - sebagai peluang optimalisasi, fokus perilaku]
- [Saran pengembangan 3 - jika perlu]
"""
    return call_openai(prompt)

def analyze_impact(impact_content):
    if impact_content is None:
        return "Analisis impact tidak dapat dilakukan karena tidak ada file impact to business yang di upload"
    
    prompt = f"""
Analisis form Impact to Business berikut dengan pendekatan APRESIATIF dan PROFESIONAL, fokus pada aspek PERILAKU:

{impact_content}

EVALUASI:
1. Perubahan PERILAKU yang terjadi dari kondisi sebelum dan sesudah implementasi program budaya
2. Peningkatan/efisiensi yang terjadi sebagai hasil dari perubahan perilaku
3. Dampak perubahan perilaku terhadap kinerja bisnis

FOKUS UTAMA - Aspek PERILAKU (BUKAN TEKNIS):
- Peningkatan kolaborasi dan kerja sama tim
- Perbaikan komunikasi dan koordinasi
- Perubahan mindset dan budaya kerja
- Peningkatan kepemimpinan dan ownership
- Peningkatan keterlibatan dan motivasi pekerja
- Penerapan nilai-nilai AKHLAK dalam praktik kerja

TONE: Apresiatif, profesional, mengakui pencapaian

Berikan output dalam format:

**Apresiasi Pencapaian:**
[Apresiasi terhadap dampak positif yang telah dicapai, soroti perubahan perilaku yang signifikan]

**Hal yang Sudah Baik:**
- [Apresiasi spesifik 1 - perubahan perilaku positif yang terukur]
- [Apresiasi spesifik 2 - perubahan perilaku positif yang terukur]
- [Apresiasi spesifik 3 - jika ada]

**Peluang Pengembangan Lebih Lanjut:**
- [Saran 1 - peluang untuk memperkuat dampak perilaku]
- [Saran 2 - peluang untuk memperkuat dampak perilaku]
- [Saran 3 - jika perlu]
"""
    return call_openai(prompt)

def analyze_evidence_comparison(skor_total, skor_benchmark_evidence, selected_hsh, selected_fungsi):
    try:
        fungsi_data = skor_total[skor_total['Fungsi'] == selected_fungsi]
        if fungsi_data.empty:
            return "Data fungsi tidak ditemukan dalam file SKOR_TOTAL_ALL"
        
        fungsi_hsh = fungsi_data.iloc[0]['HSH'] if 'HSH' in fungsi_data.columns else selected_hsh
        fungsi_hsh_normalized = normalize_hsh(fungsi_hsh)
        
        benchmark_data = skor_benchmark_evidence[
            skor_benchmark_evidence['HSH_normalized'] == fungsi_hsh_normalized
        ]
        
        if benchmark_data.empty:
            st.warning(f"⚠️ HSH '{fungsi_hsh}' tidak ditemukan exact match di benchmark. Mencoba fuzzy matching...")
            for idx, row in skor_benchmark_evidence.iterrows():
                benchmark_hsh_norm = row['HSH_normalized']
                if fungsi_hsh_normalized in benchmark_hsh_norm or benchmark_hsh_norm in fungsi_hsh_normalized:
                    benchmark_data = skor_benchmark_evidence.iloc[[idx]]
                    st.info(f"✓ Ditemukan match: '{row.iloc[0]}' untuk HSH '{fungsi_hsh}'")
                    break
        
        if benchmark_data.empty:
            st.warning(f"⚠️ Data benchmark untuk HSH '{fungsi_hsh}' tidak ditemukan. Menggunakan benchmark 'Pertamina Group' sebagai referensi.")
            benchmark_data = skor_benchmark_evidence[
                skor_benchmark_evidence['HSH_normalized'].str.contains('PERTAMINA GROUP', na=False)
            ]
            if benchmark_data.empty:
                benchmark_data = skor_benchmark_evidence.iloc[[0]]
                st.info(f"Menggunakan benchmark: '{benchmark_data.iloc[0, 0]}'")
        
        kolom_names = ['Strategi Budaya', 'Monitoring & Evaluasi', 'Sosialisasi & Partisipasi', 
                       'Pelaporan Bulanan', 'Apresiasi Pelanggan', 'Pemahaman Program', 
                       'Reward & Consequences', 'SK AoC', 'Impact to Business']
        fungsi_values = {}
        for i, name in enumerate(kolom_names):
            col_idx = 3 + i
            if col_idx < len(fungsi_data.columns):
                fungsi_values[name] = fungsi_data.iloc[0, col_idx]
        
        benchmark_values = {}
        for i, name in enumerate(kolom_names):
            col_idx = 1 + i
            if col_idx < len(benchmark_data.columns):
                benchmark_values[name] = benchmark_data.iloc[0, col_idx]
        
        differences = {}
        for name in kolom_names:
            if name in fungsi_values and name in benchmark_values:
                try:
                    diff = float(fungsi_values[name]) - float(benchmark_values[name])
                    differences[name] = diff
                except:
                    differences[name] = 'N/A'
        
        benchmark_hsh_display = benchmark_data.iloc[0, 0]
        
        comparison_text = f"""
PERBANDINGAN EVIDENCE

Fungsi: {selected_fungsi}
HSH Fungsi: {fungsi_hsh}
HSH Benchmark: {benchmark_hsh_display}

=== DATA FUNGSI ===
"""
        for name, value in fungsi_values.items():
            comparison_text += f"- {name}: {value}\n"
        
        comparison_text += f"""
=== BENCHMARK ({benchmark_hsh_display}) ===
"""
        for name, value in benchmark_values.items():
            comparison_text += f"- {name}: {value}\n"
        
        comparison_text += f"""
=== SELISIH (Fungsi - Benchmark) ===
"""
        for name, diff in differences.items():
            if diff != 'N/A':
                status = "✓ LEBIH BAIK" if diff > 0 else "⚠ PELUANG PENGEMBANGAN" if diff < 0 else "= SESUAI"
                comparison_text += f"- {name}: {diff:+.2f} {status}\n"
        
        comparison_text += """
Catatan:
- Nilai positif (+) = Fungsi LEBIH BAIK dari benchmark
- Nilai negatif (-) = Fungsi memiliki PELUANG PENGEMBANGAN
"""
        
        prompt = f"""
Analisis perbandingan Evidence berikut dengan pendekatan APRESIATIF dan PROFESIONAL:

{comparison_text}

EVALUASI:
Bandingkan performa fungsi dengan benchmark pada aspek:
1. Strategi Budaya dan implementasinya
2. Monitoring & Evaluasi oleh AoC dan Pimpinan
3. Sosialisasi & Partisipasi dalam program budaya
4. Sistem pelaporan dan apresiasi
5. Pemahaman program dan sistem reward
6. Impact to Business dari program budaya

FOKUS: Aspek PERILAKU dalam implementasi budaya kerja

TONE: Apresiatif, profesional, berbasis data

Berikan output dalam format:

**Apresiasi Pencapaian:**
[Apresiasi terhadap area yang sudah di atas atau sesuai benchmark, soroti komitmen dan konsistensi]

**Hal yang Sudah Baik:**
- [Area spesifik 1 yang di atas benchmark - dengan angka dan apresiasi]
- [Area spesifik 2 yang di atas benchmark - dengan angka dan apresiasi]
- [Area spesifik 3 - jika ada]

**Peluang Pengembangan Lebih Lanjut:**
- [Area 1 yang dapat dioptimalkan - dengan saran konkret berbasis perilaku]
- [Area 2 yang dapat dioptimalkan - dengan saran konkret berbasis perilaku]
- [Area 3 - jika perlu]
"""
        return call_openai(prompt, max_tokens=3000)
    except Exception as e:
        return f"Error dalam analisis evidence: {str(e)}\n\nDetail error: {e.__class__.__name__}"

def analyze_survei_comparison(skor_survei, skor_benchmark_survei, selected_hsh, selected_fungsi):
    try:
        fungsi_data = skor_survei[skor_survei['Fungsi'] == selected_fungsi]
        if fungsi_data.empty:
            return "Data survei fungsi tidak ditemukan dalam file Skor_SURVEI_ALL"
        
        fungsi_hsh = fungsi_data.iloc[0]['HSH'] if 'HSH' in fungsi_data.columns else selected_hsh
        fungsi_hsh_normalized = normalize_hsh(fungsi_hsh)
        
        benchmark_data = skor_benchmark_survei[
            skor_benchmark_survei['HSH_normalized'] == fungsi_hsh_normalized
        ]
        
        if benchmark_data.empty:
            st.warning(f"⚠️ HSH '{fungsi_hsh}' tidak ditemukan exact match di benchmark survei. Mencoba fuzzy matching...")
            for idx, row in skor_benchmark_survei.iterrows():
                benchmark_hsh_norm = row['HSH_normalized']
                if fungsi_hsh_normalized in benchmark_hsh_norm or benchmark_hsh_norm in fungsi_hsh_normalized:
                    benchmark_data = skor_benchmark_survei.iloc[[idx]]
                    st.info(f"✓ Ditemukan match: '{row.iloc[0]}' untuk HSH '{fungsi_hsh}'")
                    break
        
        if benchmark_data.empty:
            st.warning(f"⚠️ Data benchmark survei untuk HSH '{fungsi_hsh}' tidak ditemukan. Menggunakan benchmark 'Pertamina Group' sebagai referensi.")
            benchmark_data = skor_benchmark_survei[
                skor_benchmark_survei['HSH_normalized'].str.contains('PERTAMINA GROUP', na=False)
            ]
            if benchmark_data.empty:
                benchmark_data = skor_benchmark_survei.iloc[[0]]
                st.info(f"Menggunakan benchmark: '{benchmark_data.iloc[0, 0]}'")
        
        skor_survei_val = fungsi_data.iloc[0]['Skor Survei'] if 'Skor Survei' in fungsi_data.columns else 'N/A'
        skor_pekerja_val = fungsi_data.iloc[0]['SKOR PEKERJA'] if 'SKOR PEKERJA' in fungsi_data.columns else 'N/A'
        skor_mitra_val = fungsi_data.iloc[0]['SKOR MITRA KERJA'] if 'SKOR MITRA KERJA' in fungsi_data.columns else 'N/A'
        
        p_akhlak = fungsi_data.iloc[0]['P. AKHLAK'] if 'P. AKHLAK' in fungsi_data.columns else 'N/A'
        p_one = fungsi_data.iloc[0]['P. ONE Pertamina'] if 'P. ONE Pertamina' in fungsi_data.columns else 'N/A'
        p_program = fungsi_data.iloc[0]['P. Program Budaya'] if 'P. Program Budaya' in fungsi_data.columns else 'N/A'
        p_keberlanjutan = fungsi_data.iloc[0]['P. Keberlanjutan'] if 'P. Keberlanjutan' in fungsi_data.columns else 'N/A'
        p_safety = fungsi_data.iloc[0]['P. Safety'] if 'P. Safety' in fungsi_data.columns else 'N/A'
        
        mk_akhlak = fungsi_data.iloc[0]['MK. AKHLAK'] if 'MK. AKHLAK' in fungsi_data.columns else 'N/A'
        mk_one = fungsi_data.iloc[0]['MK. ONE Pertamina'] if 'MK. ONE Pertamina' in fungsi_data.columns else 'N/A'
        mk_program = fungsi_data.iloc[0]['MK. Program Budaya'] if 'MK. Program Budaya' in fungsi_data.columns else 'N/A'
        mk_keberlanjutan = fungsi_data.iloc[0]['MK. Keberlanjutan'] if 'MK. Keberlanjutan' in fungsi_data.columns else 'N/A'
        mk_safety = fungsi_data.iloc[0]['MK. Safety'] if 'MK. Safety' in fungsi_data.columns else 'N/A'
        
        benchmark_pekerja = benchmark_data.iloc[0, 6] if len(benchmark_data.columns) > 6 else 'N/A'
        benchmark_mitra = benchmark_data.iloc[0, 12] if len(benchmark_data.columns) > 12 else 'N/A'
        benchmark_survei = benchmark_data.iloc[0, 13] if len(benchmark_data.columns) > 13 else 'N/A'
        
        b_p_akhlak = benchmark_data.iloc[0, 1] if len(benchmark_data.columns) > 1 else 'N/A'
        b_p_one = benchmark_data.iloc[0, 2] if len(benchmark_data.columns) > 2 else 'N/A'
        b_p_program = benchmark_data.iloc[0, 3] if len(benchmark_data.columns) > 3 else 'N/A'
        b_p_keberlanjutan = benchmark_data.iloc[0, 4] if len(benchmark_data.columns) > 4 else 'N/A'
        b_p_safety = benchmark_data.iloc[0, 5] if len(benchmark_data.columns) > 5 else 'N/A'
        
        b_mk_akhlak = benchmark_data.iloc[0, 7] if len(benchmark_data.columns) > 7 else 'N/A'
        b_mk_one = benchmark_data.iloc[0, 8] if len(benchmark_data.columns) > 8 else 'N/A'
        b_mk_program = benchmark_data.iloc[0, 9] if len(benchmark_data.columns) > 9 else 'N/A'
        b_mk_keberlanjutan = benchmark_data.iloc[0, 10] if len(benchmark_data.columns) > 10 else 'N/A'
        b_mk_safety = benchmark_data.iloc[0, 11] if len(benchmark_data.columns) > 11 else 'N/A'
        
        try:
            diff_survei = float(skor_survei_val) - float(benchmark_survei) if skor_survei_val != 'N/A' and benchmark_survei != 'N/A' else 'N/A'
            diff_pekerja = float(skor_pekerja_val) - float(benchmark_pekerja) if skor_pekerja_val != 'N/A' and benchmark_pekerja != 'N/A' else 'N/A'
            diff_mitra = float(skor_mitra_val) - float(benchmark_mitra) if skor_mitra_val != 'N/A' and benchmark_mitra != 'N/A' else 'N/A'
        except:
            diff_survei = diff_pekerja = diff_mitra = 'N/A'
        
        benchmark_hsh_display = benchmark_data.iloc[0, 0]
        
        comparison_text = f"""
PERBANDINGAN SKOR SURVEI

Fungsi: {selected_fungsi}
HSH Fungsi: {fungsi_hsh}
HSH Benchmark: {benchmark_hsh_display}

=== RINGKASAN SKOR FUNGSI ===
• Skor Survei Total: {skor_survei_val}
• SKOR PEKERJA: {skor_pekerja_val}
  - P. AKHLAK: {p_akhlak}
  - P. ONE Pertamina: {p_one}
  - P. Program Budaya: {p_program}
  - P. Keberlanjutan: {p_keberlanjutan}
  - P. Safety: {p_safety}

• SKOR MITRA KERJA: {skor_mitra_val}
  - MK. AKHLAK: {mk_akhlak}
  - MK. ONE Pertamina: {mk_one}
  - MK. Program Budaya: {mk_program}
  - MK. Keberlanjutan: {mk_keberlanjutan}
  - MK. Safety: {mk_safety}

=== BENCHMARK ({benchmark_hsh_display}) ===
• Skor Survei Total: {benchmark_survei}
• SKOR PEKERJA: {benchmark_pekerja}
  - P. AKHLAK: {b_p_akhlak}
  - P. ONE Pertamina: {b_p_one}
  - P. Program Budaya: {b_p_program}
  - P. Keberlanjutan: {b_p_keberlanjutan}
  - P. Safety: {b_p_safety}

• SKOR MITRA KERJA: {benchmark_mitra}
  - MK. AKHLAK: {b_mk_akhlak}
  - MK. ONE Pertamina: {b_mk_one}
  - MK. Program Budaya: {b_mk_program}
  - MK. Keberlanjutan: {b_mk_keberlanjutan}
  - MK. Safety: {b_mk_safety}

=== SELISIH (Fungsi - Benchmark) ===
• Skor Survei Total: {diff_survei} {'✓' if diff_survei != 'N/A' and diff_survei > 0 else '⚠' if diff_survei != 'N/A' and diff_survei < 0 else ''}
• SKOR PEKERJA: {diff_pekerja} {'✓' if diff_pekerja != 'N/A' and diff_pekerja > 0 else '⚠' if diff_pekerja != 'N/A' and diff_pekerja < 0 else ''}
• SKOR MITRA KERJA: {diff_mitra} {'✓' if diff_mitra != 'N/A' and diff_mitra > 0 else '⚠' if diff_mitra != 'N/A' and diff_mitra < 0 else ''}

Catatan:
✓ = Fungsi LEBIH BAIK dari benchmark
⚠ = Fungsi memiliki PELUANG PENGEMBANGAN
"""
        
        prompt = f"""
Analisis perbandingan Survei berikut dengan pendekatan APRESIATIF dan PROFESIONAL:

{comparison_text}

EVALUASI:
Bandingkan persepsi pekerja dan mitra kerja terhadap implementasi budaya pada fungsi dengan benchmark, meliputi:
1. Pemahaman dan penerapan nilai AKHLAK
2. Implementasi ONE Pertamina
3. Partisipasi dalam Program Budaya
4. Komitmen terhadap Keberlanjutan
5. Budaya Safety

FOKUS: Aspek PERILAKU - persepsi dan pengalaman pekerja & mitra kerja terhadap budaya kerja

TONE: Apresiatif, profesional, berbasis data survei

Berikan output dalam format:

**Apresiasi Pencapaian:**
[Apresiasi terhadap skor yang sudah di atas atau sesuai benchmark, soroti area kekuatan dalam persepsi pekerja dan mitra kerja]

**Hal yang Sudah Baik:**
- [Area spesifik 1 dengan skor di atas benchmark - apresiasi dengan data]
- [Area spesifik 2 dengan skor di atas benchmark - apresiasi dengan data]
- [Area spesifik 3 - jika ada]

**Peluang Pengembangan Lebih Lanjut:**
- [Area 1 yang dapat ditingkatkan - saran konkret untuk meningkatkan persepsi dan pengalaman]
- [Area 2 yang dapat ditingkatkan - saran konkret untuk meningkatkan persepsi dan pengalaman]
- [Area 3 - jika perlu]
"""
        return call_openai(prompt, max_tokens=3500)
    except Exception as e:
        return f"Error dalam analisis survei: {str(e)}\n\nDetail error: {e.__class__.__name__}"

def create_word_document(fungsi_name, analyses):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    title = doc.add_heading('Rapport Writer Assistance', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    today = datetime.now().strftime('%d %B %Y')
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'Laporan Analisis Implementasi Budaya Kerja\n{fungsi_name}\n{today}')
    subtitle_run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        'Laporan ini disusun dengan pendekatan apresiatif untuk memberikan gambaran komprehensif '
        'mengenai implementasi budaya kerja dengan fokus pada aspek perilaku (behavior). '
        'Analisis dilakukan berdasarkan data evidence, survei, dan perbandingan dengan benchmark.'
    )
    intro_run.italic = True
    doc.add_paragraph()
    
    doc.add_heading('1. Analisis Strategi Budaya', 1)
    doc.add_paragraph(analyses['strategi_budaya'])
    doc.add_paragraph()
    
    doc.add_heading('2. Analisis Program Budaya', 1)
    doc.add_paragraph(analyses['program_budaya'])
    doc.add_paragraph()
    
    doc.add_heading('3. Analisis Impact to Business', 1)
    doc.add_paragraph(analyses['impact'])
    doc.add_paragraph()
    
    doc.add_heading('4. Analisis Perbandingan Evidence dengan Benchmark', 1)
    doc.add_paragraph(analyses['evidence_comparison'])
    doc.add_paragraph()
    
    doc.add_heading('5. Analisis Perbandingan Survei dengan Benchmark', 1)
    doc.add_paragraph(analyses['survei_comparison'])
    doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        'Laporan ini disusun sebagai bahan refleksi dan pengembangan berkelanjutan dalam implementasi '
        'budaya kerja. Kami mengapresiasi komitmen dan dedikasi seluruh tim dalam mewujudkan '
        'transformasi budaya yang positif dan berkelanjutan.'
    )
    closing_run.italic = True
    
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'\nDibuat oleh Rapport Writer Assistance\n{datetime.now().strftime("%d %B %Y, %H:%M WIB")}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Main App
def main():
    st.title("📊 Rapport Writer Assistance")
    st.caption("Asisten Analisis Implementasi Budaya Kerja dengan Pendekatan Apresiatif")
    
    with st.expander("📖 PETUNJUK PENGGUNAAN", expanded=True):
        st.markdown("""
        ### Selamat Datang di Rapport Writer Assistance!
        
        Aplikasi ini menggunakan **OpenAI AI** dan tampilan **nuansa abu-abu profesional**.
        
        **Langkah-langkah Penggunaan:**
        1. Pilih HSH dan Fungsi di sidebar
        2. Upload file PCB dan Impact (opsional)
        3. Klik **"🚀 Mulai Analisis"**
        4. Download hasil dalam format .docx
        """)

    with st.spinner('Memuat data...'):
        skor_total, skor_survei, skor_benchmark_evidence, skor_benchmark_survei = load_excel_files()
    
    if skor_total is None:
        st.stop()
    
    st.sidebar.header("⚙️ Pengaturan Analisis")
    
    hsh_list = sorted(skor_total['HSH'].unique().tolist())
    selected_hsh = st.sidebar.selectbox("Pilih HSH:", options=hsh_list)
    filtered_fungsi = sorted(skor_total[skor_total['HSH'] == selected_hsh]['Fungsi'].unique().tolist())
    selected_fungsi = st.sidebar.selectbox("Pilih Fungsi:", options=filtered_fungsi)
    
    st.sidebar.markdown("---")
    st.sidebar.subheader("📁 Upload Dokumen")
    uploaded_pcb = st.sidebar.file_uploader("Upload PCB", type=['xlsx', 'xls', 'pdf', 'png', 'jpg', 'jpeg'])
    uploaded_impact = st.sidebar.file_uploader("Upload Impact to Business", type=['xlsx', 'xls', 'pdf', 'png', 'jpg', 'jpeg'])
    st.sidebar.markdown("---")
    
    # 🔲 TOMBOL MULAI ANALISIS - NUANSA ABU-ABU
    st.markdown("""
    <style>
    div.stButton > button {
        background-color: #6c757d;
        color: white;
        border: none;
        padding: 12px 24px;
        border-radius: 8px;
        font-weight: bold;
        font-size: 16px;
        transition: background-color 0.3s ease;
        width: 100%;
    }
    div.stButton > button:hover {
        background-color: #5a6268;
    }
    </style>
    """, unsafe_allow_html=True)

    analyze_button = st.sidebar.button("🚀 Mulai Analisis", use_container_width=True)
    
    if analyze_button:
        if uploaded_pcb is None:
            st.error("⚠️ Silakan upload file PCB terlebih dahulu!")
            st.stop()
        
        st.success(f"✅ Memproses analisis untuk **{selected_fungsi}** (HSH: {selected_hsh})")
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        pcb_content = read_uploaded_file(uploaded_pcb)
        impact_content = read_uploaded_file(uploaded_impact) if uploaded_impact else None
        
        status_text.text("🔍 Menganalisis Strategi Budaya...")
        progress_bar.progress(25)
        strategi_budaya = analyze_strategi_budaya(pcb_content)
        
        status_text.text("🔍 Menganalisis Program Budaya...")
        progress_bar.progress(40)
        program_budaya = analyze_program_budaya(pcb_content)
        
        status_text.text("🔍 Menganalisis Impact to Business...")
        progress_bar.progress(55)
        impact = analyze_impact(impact_content)
        
        status_text.text("🔍 Menganalisis Perbandingan Evidence...")
        progress_bar.progress(70)
        evidence_comparison = analyze_evidence_comparison(skor_total, skor_benchmark_evidence, selected_hsh, selected_fungsi)
        
        status_text.text("🔍 Menganalisis Perbandingan Survei...")
        progress_bar.progress(85)
        survei_comparison = analyze_survei_comparison(skor_survei, skor_benchmark_survei, selected_hsh, selected_fungsi)
        
        analyses = {
            'strategi_budaya': strategi_budaya,
            'program_budaya': program_budaya,
            'impact': impact,
            'evidence_comparison': evidence_comparison,
            'survei_comparison': survei_comparison
        }
        
        status_text.text("📝 Membuat dokumen Word...")
        progress_bar.progress(95)
        doc_io = create_word_document(selected_fungsi, analyses)
        
        progress_bar.progress(100)
        status_text.text("✅ Analisis selesai!")
        st.balloons()
        
        st.markdown("---")
        st.header("📊 Hasil Analisis")
        
        # 🔲 TAB - NUANSA ABU-ABU
        st.markdown("""
        <style>
        .stTabs [data-baseweb="tab-list"] {
            background-color: #f8f9fa;
            padding: 10px;
            border-radius: 8px;
        }
        .stTabs [data-baseweb="tab"] {
            height: 40px;
            white-space: pre-wrap;
            background-color: #e9ecef;
            border-radius: 6px;
            color: #495057;
            font-weight: bold;
            padding: 0 16px;
            margin-right: 8px;
        }
        .stTabs [aria-selected="true"] {
            background-color: #6c757d;
            color: white;
        }
        </style>
        """, unsafe_allow_html=True)

        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "Strategi Budaya", 
            "Program Budaya", 
            "Impact to Business",
            "Perbandingan Evidence",
            "Perbandingan Survei"
        ])
        
        with tab1: st.markdown("### Analisis Strategi Budaya\n" + strategi_budaya)
        with tab2: st.markdown("### Analisis Program Budaya\n" + program_budaya)
        with tab3: st.markdown("### Analisis Impact to Business\n" + impact)
        with tab4: st.markdown("### Analisis Perbandingan Evidence\n" + evidence_comparison)
        with tab5: st.markdown("### Analisis Perbandingan Survei\n" + survei_comparison)
        
        st.markdown("---")
        today = datetime.now().strftime('%m_%d')
        filename = f"Rapp_{selected_fungsi.replace(' ', '_').replace('/', '_')}_{today}.docx"

        # 🔲 TOMBOL DOWNLOAD - NUANSA ABU-ABU
        st.markdown("""
        <style>
        .stDownloadButton > button {
            background-color: #6c757d !important;
            color: white !important;
            border: none !important;
            padding: 12px 24px !important;
            border-radius: 8px !important;
            font-weight: bold !important;
            font-size: 16px !important;
            width: 100% !important;
            transition: background-color 0.3s ease !important;
        }
        .stDownloadButton > button:hover {
            background-color: #5a6268 !important;
        }
        </style>
        """, unsafe_allow_html=True)

        st.download_button(
            label="📥 Download Hasil Analisis (.docx)",
            data=doc_io,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        # 🔲 PESAN SUKSES - NUANSA ABU-ABU MUDA
        st.markdown(f"""
        <div style="
            background-color: #f8f9fa;
            padding: 12px;
            border-radius: 8px;
            border-left: 4px solid #6c757d;
            margin-top: 10px;
            color: #495057;
            font-weight: bold;
        ">
            ✅ Dokumen siap didownload: <strong>{filename}</strong>
        </div>
        """, unsafe_allow_html=True)
    
    else:
        st.info("👈 Silakan pilih HSH, Fungsi, upload file, dan klik tombol **Mulai Analisis** di sidebar")
        col1, col2 = st.columns(2)
        with col1: st.metric("HSH Terpilih", selected_hsh if selected_hsh else "-")
        with col2: st.metric("Fungsi Terpilih", selected_fungsi if selected_fungsi else "-")
        
        st.markdown("---")
        st.markdown("### 💡 Tips Penggunaan")
        st.markdown("""
        - Semua analisis menggunakan **pendekatan apresiatif**
        - Fokus pada **perubahan perilaku**, bukan teknis
        - API key disimpan aman melalui **Streamlit Secrets**
        """)

if __name__ == "__main__":

    main()

