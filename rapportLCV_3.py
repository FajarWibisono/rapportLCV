import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from PIL import Image
import io
from datetime import datetime
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from groq import Groq

def initialize_groq_client():
    """Initialize Groq client with API key from secrets"""
    try:
        # Use secrets for Streamlit Cloud deployment
        api_key = st.secrets["groq"]["api_key"]
        return Groq(api_key=api_key)
    except KeyError:
        st.error("""
        🔑 **API Key Missing**  
        Please add your Groq API key to the secrets configuration:
        ```
        [groq]
        api_key = "your_api_key_here"
        ```
        """)
        st.stop()
    except Exception as e:
        st.error(f"❌ Failed to initialize Groq client: {str(e)}")
        st.stop()

# Initialize Groq client
client = initialize_groq_client()

def read_pdf_text(file):
    """Membaca teks dari file PDF dengan error handling yang robust"""
    try:
        if not hasattr(file, 'read'):
            return ""
        
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error membaca PDF: {str(e)}")
        return ""

def read_excel_data(file):
    """Membaca data dari file Excel dengan validasi yang kuat"""
    try:
        if not hasattr(file, 'read'):
            return None
        
        # Try different engine options for compatibility
        try:
            df = pd.read_excel(file, engine='openpyxl')
        except:
            try:
                df = pd.read_excel(file, engine='xlrd')
            except:
                df = pd.read_excel(file)
        return df
    except Exception as e:
        st.error(f"Error membaca Excel: {str(e)}")
        return None

def analyze_with_groq(prompt, max_tokens=500):
    """Fungsi untuk mengirimkan prompt ke model Groq dengan error handling lengkap"""
    try:
        with st.spinner("🧠 Processing with AI..."):
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="moonshotai/kimi-k2-instruct-0905",
                max_tokens=max_tokens,
                temperature=0.3,
            )
            return chat_completion.choices[0].message.content.strip()
    except Exception as e:
        error_msg = f"""
        ❌ Error saat memanggil API Groq:
        {str(e)}
        
        **Possible solutions:**
        1. Check your internet connection
        2. Verify your API key is valid
        3. Ensure you have sufficient quota
        4. Try again in a few minutes
        
        **Temporary workaround:**
        Using fallback analysis until connection is restored.
        """
        st.error(error_msg)
        # Return fallback analysis
        return "Analisis sementara: Data telah diproses namun terjadi kendala koneksi dengan AI. Silakan coba lagi nanti atau hubungi administrator."

def analyze_pcb_quality(pcb_text):
    """Analisis kualitas PCB dengan pendekatan positif"""
    if not pcb_text or not pcb_text.strip():
        return "Teks PCB tidak tersedia untuk dianalisis. Pastikan file PCB dapat dibaca dengan baik."
    
    prompt = f"""
    Berikan analisis yang positif dan memotivasi terhadap isi dokumen Program Budaya (PCB) berikut:
    "{pcb_text[:2000]}"  # Limit text to avoid token overflow
    
    Fokuskan pada:
    - Apakah Goals Business Initiatives/improvement dirumuskan dengan jelas dan menggunakan prinsip SMART?
    - Apakah alur logika antara identifikasi peluang bisnis dan inisiatif perbaikan terasa runut dan kuat?
    - Apa saja kekuatan utama dari pendekatan strategis yang digunakan?
    
    Gunakan nada apresiatif, optimis, dan memberi semangat.
    Di akhir, tambahkan bagian:
    - **Hal yang Dapat Ditingkatkan** untuk improvement di tahun berikutnya.
    
    Maksimal 270 kata.
    """
    return analyze_with_groq(prompt)

def analyze_cultural_programs(pcb_text):
    """Analisis Program Budaya dengan fokus pada perubahan perilaku"""
    if not pcb_text or not pcb_text.strip():
        return "Teks PCB tidak tersedia untuk dianalisis. Pastikan file PCB dapat dibaca dengan baik."
    
    prompt = f"""
    Analisis Program Budaya dalam dokumen berikut:
    "{pcb_text[:2000]}"  # Limit text to avoid token overflow
    
    Evaluasi:
    - Program Standar (One Hour Meeting)
    - Program Mandatory (ONE Action)
    - Program Spesifik (ONE KOLAB)
    
    Untuk masing-masing program:
    - Apakah judul dan deliverables-nya mencerminkan tujuan dengan baik?
    - Bagaimana kontribusi program-program ini terhadap pencapaian Goals Business Initiatives?
    - Berikan apresiasi terhadap hal-hal yang sudah kuat.
    - Fokus pada perubahan perilaku yang terjadi dan yang masih diperlukan
    
    Tulis bagian "Hal yang Dapat Ditingkatkan" dengan fokus pada perubahan perilaku konkret yang diharapkan.
    Gunakan nada positif namun tetap memberikan insight pengembangan.
    Maksimal 270 kata.
    """
    return analyze_with_groq(prompt)

def analyze_business_impact(impact_text):
    """Analisis Impact to Business dengan struktur yang jelas"""
    if not impact_text or not impact_text.strip():
        return "Data impact tidak tersedia untuk dianalisis. Pastikan file Kuantifikasi Impact dapat dibaca dengan baik."
    
    prompt = f"""
    Analisis form impact to business berikut:
    "{impact_text[:2000]}"  # Limit text to avoid token overflow
    
    Jawab dengan struktur:
    1. **Perubahan yang Terjadi**: jelaskan transformasi dari sebelum ke sesudah.
    2. **Efisiensi/Peningkatan**: sebutkan angka, waktu, biaya, atau produktivitas yang meningkat.
    3. **Hal yang Sudah Baik**: soroti keberhasilan dan praktik unggulan.
    4. **Hal yang Dapat Ditingkatkan**: sampaikan saran pengembangan secara konstruktif.
    
    Gunakan bahasa profesional, positif, dan membangun.
    Pastikan total jawaban tidak lebih dari 360 kata.
    """
    return analyze_with_groq(prompt, max_tokens=600)

def analyze_lcv_evidence(evidence_df, current_lcv_score):
    """Analisis Evidence Implementasi LCV"""
    if evidence_df is None or evidence_df.empty:
        return "Data evidence tidak tersedia untuk dianalisis. Pastikan file Evidence dapat dibaca dengan baik."
    
    try:
        # Get the last row (most recent data)
        last_row = evidence_df.iloc[-1]
        evidence_text = f"Data terakhir:\n{last_row.to_string()}"
        
        # Get column names related to notes
        note_columns = [col for col in evidence_df.columns if 'catatan' in str(col).lower() or 'note' in str(col).lower()]
        
        # Prepare specific columns for analysis
        monev_val = ""
        sosialisasi_val = ""
        pelaporan_val = ""
        reward_val = ""
        
        # Try to find relevant columns with multiple naming possibilities
        for col in evidence_df.columns:
            col_lower = str(col).lower()
            if 'monev' in col_lower or 'monitoring' in col_lower or 'evaluasi' in col_lower:
                monev_val = str(last_row[col]) if col in last_row else ""
            if 'sosialisasi' in col_lower or 'partisipasi' in col_lower:
                sosialisasi_val = str(last_row[col]) if col in last_row else ""
            if 'lapor' in col_lower or 'report' in col_lower:
                pelaporan_val = str(last_row[col]) if col in last_row else ""
            if 'reward' in col_lower or 'consequence' in col_lower:
                reward_val = str(last_row[col]) if col in last_row else ""
        
        prompt = f"""
        Analisis implementasi LCV berdasarkan data evidence terbaru.
        Skor akhir LCV: {current_lcv_score}

        Fokus pada:
        - **Skor Akhir**: apresiasi terhadap pencapaian.
        - **Catatan Monev oleh AoC dan Pimpinan**: "{monev_val}"
        - **Catatan Sosialisasi dan Partisipasi**: "{sosialisasi_val}"
        - **Pelaporan Bulanan**: apakah konsisten dan tepat waktu? "{pelaporan_val}"
        - **Reward/Consequences**: apakah sudah dilakukan? "{reward_val}"

        Berikan analisis yang membangun dan positif. Soroti keberhasilan, dan sampaikan juga hal-hal yang bisa ditingkatkan.
        Untuk catatan yang berisi "belum ada update" atau "tidak ada perubahan", anggap sebagai indikator positif bahwa aktivitas berjalan lancar.
        Maksimal 270 kata.
        """
        return analyze_with_groq(prompt)
    except Exception as e:
        st.error(f"Error processing evidence data: {str(e)}")
        return "Gagal menganalisis data evidence. Silakan periksa format file Evidence Anda."

def sanitize_filename(name):
    """Menghapus karakter ilegal dari nama file"""
    name = re.sub(r'[^\w\-_]', '_', name)
    name = re.sub(r'_+', '_', name)
    return name.strip('_')[:50]  # Limit filename length

def create_word_document(pcb_analysis, cultural_analysis, impact_analysis, lcv_analysis, lcv_score, timestamp):
    """Membuat dokumen Word dengan error handling"""
    try:
        doc = Document()
        
        # Add title with center alignment
        title = doc.add_heading("RAPPORT WRITER ASSISTANCE - ANALYSIS REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        doc.add_heading(f"Generated on: {timestamp}", level=1)
        
        # Add sections
        doc.add_heading("1. Analisis Kualitas PCB", level=2)
        doc.add_paragraph(pcb_analysis)
        
        doc.add_heading("2. Analisis Program Budaya", level=2)
        doc.add_paragraph(cultural_analysis)
        
        doc.add_heading("3. Analisis Impact to Business", level=2)
        doc.add_paragraph(impact_analysis)
        
        doc.add_heading("4. Analisis Evidence Implementasi LCV", level=2)
        doc.add_paragraph(lcv_analysis)
        
        # Add footer
        doc.add_paragraph("\n" + "-"*50)
        doc.add_paragraph("Generated by RAPPORT WRITER ASSISTANCE - Streamlit App")
        doc.add_paragraph(f"LCV Score: {lcv_score}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

# Page configuration
st.set_page_config(
    page_title="RAPPORT WRITER ASSISTANCE",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .reportview-container {
        margin-top: -3em;
    }
    .stProgress > div > div > div > div {
        background-color: #1f77b4;
    }
    .st-bb {
        background-color: transparent;
    }
    .st-at {
        background-color: transparent;
    }
    footer {
        font-size: 14px;
    }
</style>
""", unsafe_allow_html=True)

# Main title with emoji
st.title("📝 RAPPORT WRITER ASSISTANCE")

# Instructions
with st.expander("ℹ️ Petunjuk Penggunaan", expanded=False):
    st.markdown("""
    **Langkah-langkah penggunaan:**
    1. Upload skor LCV tahun lalu (angka antara 0-500)
    2. Upload file PCB (PDF/JPEG) 
    3. Upload file Kuantifikasi Impact (PDF/Excel)
    4. Upload file Evidence (Excel)
    5. Klik tombol 'Generate Analysis' untuk mendapatkan hasil analisis
    6. Klik tombol 'Export to Word' untuk mengunduh laporan
    
    **Catatan:**
    - Pastikan file yang diupload sesuai format yang diminta
    - Koneksi internet diperlukan untuk analisis dengan AI
    - Proses analisis mungkin memakan waktu beberapa detik
    """)

# Sidebar configuration
with st.sidebar:
    st.header("📁 Upload Dokumen")
    
    # LCV Score input
    lcv_score = st.number_input(
        "Skor LCV Tahun Lalu (0-500)", 
        min_value=0, 
        max_value=500, 
        value=0,
        help="Masukkan skor LCV tahun lalu dalam angka antara 0-500"
    )
    
    # File uploads with clear instructions
    pcb_file = st.file_uploader(
        "Upload File PCB (PDF/JPEG)", 
        type=['pdf', 'jpeg', 'jpg', 'png'],
        help="Upload file PCB dalam format PDF, JPEG, atau PNG"
    )
    
    impact_file = st.file_uploader(
        "Upload File Kuantifikasi Impact (PDF/Excel)", 
        type=['pdf', 'xlsx', 'xls'],
        help="Upload file Kuantifikasi Impact dalam format PDF atau Excel"
    )
    
    evidence_file = st.file_uploader(
        "Upload File Evidence (Excel)", 
        type=['xlsx', 'xls'],
        help="Upload file Evidence dalam format Excel"
    )
    
    # Analysis button with primary style
    analyze_button = st.button("✨ Generate Analysis", type="primary", use_container_width=True)
    
    # Add footer to sidebar
    st.markdown("---")
    st.caption("© 2025 RAPPORT WRITER ASSISTANCE")
    st.caption("Powered by Groq AI & Streamlit")

# Main content area
if analyze_button:
    if not (lcv_score and pcb_file and impact_file and evidence_file):
        st.warning("⚠️ Mohon lengkapi semua field dan upload semua file yang diperlukan!")
        st.stop()
    
    try:
        st.header("📊 Hasil Analisis")
        st.caption("Hasil analisis dihasilkan menggunakan AI model moonshotai/kimi-k2-instruct-0905")
        
        # Process PCB file
        with st.spinner("Processing PCB file..."):
            pcb_text = ""
            if pcb_file.name.lower().endswith('.pdf'):
                pcb_text = read_pdf_text(pcb_file)
            else:
                # Display image but don't use it for analysis
                image = Image.open(pcb_file)
                with st.expander("📸 Preview File PCB (Gambar)", expanded=False):
                    st.image(image, caption="File PCB (Gambar)", use_column_width=True)
                pcb_text = "File PCB berupa gambar - analisis teks menggunakan deskripsi umum."
        
        # Process Impact file
        with st.spinner("Processing Kuantifikasi Impact file..."):
            impact_text = ""
            if impact_file.name.lower().endswith('.pdf'):
                impact_text = read_pdf_text(impact_file)
            else:
                impact_df = read_excel_data(impact_file)
                if impact_df is not None:
                    impact_text = impact_df.to_string()
                else:
                    impact_text = "Data Kuantifikasi Impact tidak dapat dibaca."
        
        # Process Evidence file
        with st.spinner("Processing Evidence file..."):
            evidence_df = read_excel_data(evidence_file)
        
        # Generate analyses
        st.subheader("🔍 Analisis Kualitas PCB")
        pcb_analysis = analyze_pcb_quality(pcb_text)
        st.write(pcb_analysis)
        
        st.subheader("🎯 Analisis Program Budaya")
        cultural_analysis = analyze_cultural_programs(pcb_text)
        st.write(cultural_analysis)
        
        st.subheader("📈 Analisis Impact to Business")
        impact_analysis = analyze_business_impact(impact_text)
        st.write(impact_analysis)
        
        st.subheader("📋 Analisis Evidence Implementasi LCV")
        lcv_analysis = analyze_lcv_evidence(evidence_df, lcv_score)
        st.write(lcv_analysis)
        
        # Export to Word
        st.header("💾 Export Hasil Analisis")
        
        # Create timestamp for filename
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        filename_timestamp = datetime.now().strftime("%y_%m_%d_%H-%M-%S")
        
        # Create Word document
        word_buffer = create_word_document(
            pcb_analysis, 
            cultural_analysis, 
            impact_analysis, 
            lcv_analysis, 
            lcv_score, 
            timestamp
        )
        
        if word_buffer:
            # Generate sanitized filename
            default_filename = f"Rap_General_{filename_timestamp}.docx"
            
            # Download button
            st.download_button(
                label="📥 Download Laporan (Word)",
                data=word_buffer,
                file_name=default_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
            st.success("✅ Laporan berhasil dibuat! Klik tombol di atas untuk mengunduh.")
        else:
            st.error("❌ Gagal membuat dokumen Word. Silakan coba lagi.")
    
    except Exception as e:
        st.error(f"""
        ❌ Terjadi kesalahan tidak terduga:
        {str(e)}
        
        **Silakan coba langkah berikut:**
        1. Refresh halaman
        2. Periksa kembali file yang diupload
        3. Pastikan koneksi internet stabil
        4. Hubungi administrator jika masalah berlanjut
        """)
        st.exception(e)  # Only for debugging, remove in production
else:
    st.info("👆 Silakan upload dokumen di sidebar dan klik 'Generate Analysis' untuk memulai analisis.")
    st.image("https://streamlit.io/images/brand/streamlit-logo-secondary-colormark-darktext.png", width=200)

# Footer
st.markdown("---")
st.caption("RAPPORT WRITER ASSISTANCE - AI Powered Report Analysis Tool | © 2025")